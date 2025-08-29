# bookgen/main.py
import os, json, re, sys
from datetime import datetime
from pathlib import Path
import yaml

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openai import OpenAI

# ====================== CONFIG ======================
MODEL                = os.getenv("BOOK_MODEL", "gpt-4o")
SUBSECTION_TOKENS    = int(os.getenv("SUBSECTION_MAX_OUTPUT_TOKENS", "7000"))
CHAPTER_TOKENS       = int(os.getenv("CHAPTER_MAX_OUTPUT_TOKENS", "18000"))
MIN_SUBSECTION_WORDS = 450          # target acceptance
MIN_HARD_WORDS       = 200          # below this -> ALWAYS retry
MAX_TRIES_PER_SUB    = 5            # more chances to avoid empty/short sections

# Mini-headings normalization style: "bullet" or "bold"
MINI_MODE = os.getenv("MINI_HEADING_MODE", "bullet").strip().lower()

DOCS_DIR   = Path("output")
CHECKPOINT = Path("progress.json")
BOOK_YAML  = Path("book.yaml")

# Always create a NEW file
RUN_ID = os.getenv("BOOK_RUN_ID") or datetime.now().strftime("%Y%m%d-%H%M%S")

# ===== MASTER PROMPT (global, reused for all books) =====
MASTER_PROMPT = """
You are a book-writing assistant.  
Your job is to generate high-quality book content that follows the provided Table of Contents.  

Rules:  
1) Follow the TOC exactly; do not invent extra chapters or headings.  
2) For each subheading, write ~500–600 words of complete, on-topic content.  
3) Avoid repetition; each sentence must add new value.  
4) Stay strictly on-topic for each subheading.  
5) Ensure smooth flow and coherence across paragraphs.  
6) Never use placeholders or meta comments.  
7) No decorative separators (—, ***, etc.).  
8) Do not restate “Chapter/Day …” or the subheading line inside the body.  

Global quality:  
- Each section must feel complete, rich, and useful.  
- Maintain consistent length and depth across the entire book.  
- The output must be copy-paste ready.  

# ====================================================

client = OpenAI()

# ---------------- File IO ----------------
def load_yaml(path=BOOK_YAML):
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def _set_mirror_margins(doc: Document):
    """Turn on mirror margins in document settings (best-effort; OK if it fails)."""
    try:
        settings_part = doc._part.package.settings_part
        s = settings_part.element
        node = s.find(qn('w:mirrorMargins'))
        if node is None:
            s.append(OxmlElement('w:mirrorMargins'))
    except Exception:
        pass  # not critical

def ensure_doc(title: str):
    """Always create a NEW .docx with RUN_ID and set page/margin/spacing defaults."""
    DOCS_DIR.mkdir(exist_ok=True)
    safe_title = re.sub(r'[\\/:*?"<>|]+', '-', title).strip()
    doc_path = DOCS_DIR / f"BOOK - {safe_title} - {RUN_ID}.docx"

    doc = Document()

    # ---- Page size & margins ----
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)
    try:
        section.gutter = Inches(0.2)
    except Exception:
        pass
    _set_mirror_margins(doc)

    # ---- Base style (Normal) ----
    normal = doc.styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(13)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)  # small visual gap
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.2

    # Apply same spacing rules to headings so Word doesn't add extra gaps
    for style_name in ("Heading 1", "Heading 2", "Title"):
        try:
            st = doc.styles[style_name]
            st.font.name = "Cambria"
            st.font.size = Pt(18 if style_name == "Title" else 16)
            sp = st.paragraph_format
            sp.space_before = Pt(0)
            sp.space_after  = Pt(0)
            sp.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            sp.line_spacing = 1.2
        except KeyError:
            pass

    h0 = doc.add_heading(title, level=0)
    h0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(str(doc_path))
    return doc, doc_path

def save_doc(doc: Document, path: Path):
    doc.save(str(path))

# ---------------- TOC helpers ----------------
def flatten_toc(toc_list):
    chapters = []
    for item in toc_list:
        if isinstance(item, str):
            chapters.append({"title": item, "subs": []})
        elif isinstance(item, dict):
            for k, v in item.items():
                subs = v if isinstance(v, list) else []
                chapters.append({"title": k, "subs": subs})
        else:
            raise ValueError("TOC item non valido: usa stringa o dict {chapter: [subs...]}")
    return chapters

# ---------------- OpenAI helpers ----------------
def responses_text(resp):
    if hasattr(resp, "output") and resp.output:
        out = []
        for block in resp.output:
            if getattr(block, "type", "") == "message":
                for c in getattr(block, "content", []) or []:
                    if getattr(c, "type", "") == "output_text":
                        out.append(c.text)
            elif getattr(block, "type", "") == "output_text":
                out.append(block.text)
        txt = "".join(out).strip()
        if txt:
            return txt
    if getattr(resp, "output_text", None):
        return resp.output_text.strip()
    if getattr(resp, "choices", None):
        return resp.choices[0].message.content.strip()
    return ""

def call_openai(prompt: str, max_tokens: int) -> str:
    r = client.responses.create(
        model=MODEL,
        input=prompt,
        max_output_tokens=max_tokens
    )
    return responses_text(r)

# ---------------- Memory (summary + claims + angle) ----------------
def build_angle(chapter: str, sub: str, rolling_summary: str, claims: list):
    clist = "\n".join(f"- {c}" for c in (claims or [])[:7]) or "- (none)"
    prompt = f"""
You are a content planner.
Given the rolling summary and covered claims, provide ONE sentence (max 25 words) that states the UNIQUE ANGLE for the next section.

CHAPTER: {chapter}
SUBHEADING: {sub}

ROLLING SUMMARY:
{rolling_summary or '(none)'}

COVERED CLAIMS:
{clist}

Rules:
- Be concrete and specific (what new example, case, or perspective to add).
- Do NOT restate the subheading or claims; propose the new angle only.
- Output ONE sentence, nothing else."""
    out = call_openai(prompt, 400) or ""
    return out.splitlines()[0].strip()

def update_memory_from_text(text: str, mem: dict):
    prompt = f"""
Summarize the following section in 120–180 words (plain, non-marketing), then list 5 KEY CLAIMS as standalone sentences.

SECTION:
{text}

Format:
SUMMARY:
<one paragraph 120–180 words>

CLAIMS:
- <sentence 1>
- <sentence 2>
- <sentence 3>
- <sentence 4>
- <sentence 5>
"""
    out = call_openai(prompt, 1200) or ""
    summary, claims = "", []
    m = re.search(r"SUMMARY:\s*([\s\S]*?)\n\s*CLAIMS:", out, re.I)
    if m: summary = m.group(1).strip()
    tail = re.split(r"CLAIMS:\s*", out, flags=re.I)
    if len(tail) > 1:
        claims = [re.sub(r"^-+\s*", "", ln).strip()
                  for ln in tail[1].splitlines() if ln.strip()]
    mem["summary"] = summary or mem.get("summary", "")
    mem["claims"]  = (claims or mem.get("claims", []))[:7]
    return mem

# ---------------- Prompt builders ----------------
END_MARK = "<<<END_OF_SUBHEADING>>>"

def subheading_prompt(master, persona, title, chapters_list, chapter, sub, mem):
    claims_list = "\n".join(f"- {c}" for c in (mem.get("claims") or [])[:7]) or "- (none)"
    return f"""
You are an expert ghostwriter.

--- MASTER PROMPT (DO NOT OUTPUT) ---
{master}

--- BUYER PERSONA & STYLE (DO NOT OUTPUT) ---
{persona}

CONTEXT SUMMARY (DO NOT OUTPUT; do not repeat verbatim):
{mem.get("summary") or '(none)'}

ALREADY COVERED CLAIMS (avoid repeating; if similar theme, add a new angle):
{claims_list}

ANGLE TO ADOPT (one sentence):
{mem.get("angle") or 'Offer a concrete, fresh angle with specific examples.'}

STYLE CLAMP (MANDATORY):
- Follow the TOC exactly; never invent new headings.
- No placeholders or meta talk.
- No Markdown headings in the body.
- Bold only short key phrases (≤ 8 words).
- Write continuous prose, not broken into many headings.
- Each subheading MUST produce ~700–800 words of prose.
- Never skip or merge subheadings, even if they overlap; bring a NEW example or angle.

--- BOOK CONTEXT ---
Title: {title}

GLOBAL CHAPTER LIST:
{chapters_list}

CURRENT CHAPTER: {chapter}
CURRENT SUBHEADING: {sub}

OUTPUT:
- Write ~500–600 words of rich, specific content for THIS subheading only.
- End cleanly with prose, then on a new line output exactly: {END_MARK}
"""

def chapter_only_prompt(master, persona, title, chapters_list, chapter):
    return f"""
You are an expert ghostwriter.

--- MASTER PROMPT (DO NOT OUTPUT) ---
{master}

--- BUYER PERSONA & STYLE (DO NOT OUTPUT) ---
{persona}

STYLE CLAMP (MANDATORY):
- Warm, simple, accessible; not academic.
- No "Chapter X:" lines inside the body, no placeholders, no invented headings.
- Bold only short phrases (≤ 8 words), never whole paragraphs.
- Mini-topics allowed as bold lead-ins or compact bullets only if they truly add clarity.
The text should feel like a flowing conversation. Avoid over-fragmenting ideas into subheadings.

--- BOOK CONTEXT ---
Title: {title}

GLOBAL CHAPTER LIST:
{chapters_list}

CURRENT CHAPTER (no subheadings): {chapter}

OUTPUT:
- Write ~2,000–3,000 words of continuous, engaging prose.
- End cleanly with prose, then on a new line output exactly: {END_MARK}
"""

# ---------------- Output cleaning/formatting ----------------
BULLET_RE = re.compile(r"^(\s*)[-*•]\s+", re.MULTILINE)

def clean_text(raw: str) -> str:
    if not raw:
        return ""
    text = raw.replace(END_MARK, "").strip()
    text = re.sub(r"[\U00010000-\U0010FFFF]", "", text)                     # strip emoji
    text = re.sub(r"^#{2,3}\s+", "", text, flags=re.MULTILINE)              # strip markdown heads
    text = BULLET_RE.sub(r"\1", text)                                       # lists -> paragraphs
    # strip all-caps lines and unwrap full-bold lines
    lines = []
    for ln in text.splitlines():
        t = ln.rstrip()
        letters = re.sub(r"[^A-Za-z]+", "", t)
        upper_ratio = (sum(1 for ch in letters if ch.isupper()) / len(letters)) if letters else 0
        if upper_ratio > 0.6 and len(t) > 8:
            continue
        if re.match(r"^\*\*[^*].*[^*]\*\*$", t):
            t = re.sub(r"^\*\*(.*)\*\*$", r"\1", t)
        lines.append(t)
    return "\n".join(lines).strip()

def is_probable_mini_heading(line: str) -> bool:
    t = line.strip()
    if not t or len(t.split()) < 2 or len(t.split()) > 7: return False
    if re.search(r"[.!?]$", t): return False
    if re.match(r"^(chapter|day)\s+\d+:", t, flags=re.I): return False
    if re.match(r"^\d", t): return False
    words = t.split()
    caps = sum(1 for w in words if re.match(r"^[A-Z][a-z]+$", w))
    return caps / len(words) >= 0.6

def normalize_mini_headings(text: str) -> str:
    lines = [ln for ln in text.splitlines()]
    out = []
    i = 0
    while i < len(lines):
        cur = lines[i].strip()
        nxt = lines[i+1].strip() if i+1 < len(lines) else ""
        if is_probable_mini_heading(cur) and nxt and not is_probable_mini_heading(nxt):
            if MINI_MODE == "bold":
                merged = f"**{cur}.** {nxt}"
            else:
                merged = f"• **{cur}** — {nxt}"
            out.append(merged)
            i += 2
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)

def split_into_paragraphs_preserving_bold(line: str):
    parts = re.split(r"(\*\*)", line)
    result_spans, bold = [], False
    for part in parts:
        if part == "**":
            bold = not bold
            continue
        if not part:
            continue
        words = part.strip().split()
        if bold and len(words) > 8:
            result_spans.append((False, part))
        else:
            result_spans.append((bold, part))
    return result_spans

def add_paragraph_with_bold(doc: Document, text: str):
    p = doc.add_paragraph()
    # Paragraph formatting (Cambria 13, justified, small gap, Multiple 1.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(6)  # tiny space between paragraphs
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.2

    spans = split_into_paragraphs_preserving_bold(text)
    for is_bold, chunk in spans:
        run = p.add_run(chunk)
        run.font.name = "Cambria"
        run.font.size = Pt(13)
        run.bold = bool(is_bold)

def write_subsection(doc: Document, body_text: str):
    body_text = normalize_mini_headings(body_text)

    for para in body_text.split("\n"):
        t = para.strip()
        if not t:
            continue  # no blank paragraphs (prevents big gaps)
        # Split very long paragraphs on sentence boundaries
        if len(t.split()) > 180:
            for chunk in re.split(r"(?<=[.!?])\s+", t):
                if chunk.strip():
                    add_paragraph_with_bold(doc, chunk.strip())
        else:
            add_paragraph_with_bold(doc, t)

# ---------------- Light fixer ----------------
def quick_validate_and_fix(text: str, covered_claims: list, angle: str):
    issues = []
    if re.search(r"^#{2,3}\s+", text, flags=re.M):
        issues.append("Markdown headings inside body.")
    if re.search(r"^[A-Z0-9][A-Z0-9\s.,;:!?\"'()\\-]{20,}$", text, flags=re.M):
        issues.append("All-caps paragraph detected.")
    for c in (covered_claims or []):
        if c and len(c) > 8 and c.lower() in text.lower():
            issues.append("Repeats a previously stated claim too literally.")

    if not issues:
        return text

    prompt = f"""
You are a copy editor. Fix the issues MINIMALLY without changing meaning or tone.

ISSUES:
{os.linesep.join('- ' + i for i in issues)}

ANGLE TO KEEP:
{angle or '(keep focus; no new topics)'}

TEXT:
{text}

RULES:
- Keep length similar; do NOT shorten significantly.
- No headings, no all-caps.
- Start with prose (no repeated subheading).
- You may turn tiny standalone headings into bold lead-ins.
- Return ONLY the cleaned text, nothing else."""
    fixed = call_openai(prompt, 1800) or text
    return clean_text(fixed)

# ---------------- MAIN ----------------
def main():
    # New run => clear any previous checkpoint (we create a fresh file every time)
    if CHECKPOINT.exists():
        CHECKPOINT.unlink()

    if not BOOK_YAML.exists():
        print(f"Errore: manca {BOOK_YAML.resolve()}")
        sys.exit(1)

    cfg = load_yaml(BOOK_YAML)
    title   = cfg["title"]
    persona = cfg["persona"]
    chapters = flatten_toc(cfg["toc"])

    doc, doc_path = ensure_doc(title)

    # Memory state
    mem = {"summary": "", "claims": []}

    chapters_list = "\n".join(f"- {c['title']}" for c in chapters)

    for ch in chapters:
        h = doc.add_heading(ch["title"], level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if not ch["subs"]:
            prompt = chapter_only_prompt(MASTER_PROMPT, persona, title, chapters_list, ch["title"])
            raw = call_openai(prompt, CHAPTER_TOKENS)
            cleaned = clean_text(raw)
            fixed = quick_validate_and_fix(cleaned, mem.get("claims"), angle=None)
            write_subsection(doc, fixed)
            mem = update_memory_from_text(fixed, mem)
            save_doc(doc, doc_path)
            continue

        for sub in ch["subs"]:
            sh = doc.add_heading(sub, level=2)
            sh.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            mem["angle"] = build_angle(ch["title"], sub, mem.get("summary"), mem.get("claims"))

            tries, out_text = 0, ""
            while tries < MAX_TRIES_PER_SUB:
                prompt = subheading_prompt(MASTER_PROMPT, persona, title, chapters_list, ch["title"], sub, mem)
                raw = call_openai(prompt, SUBSECTION_TOKENS)
                cleaned = clean_text(raw)

                # ---- Looser echo filter: only drop echoed headings if we ALSO have real prose
                n_ch  = ch["title"].strip().lower()
                n_sub = sub.strip().lower()
                lines = cleaned.splitlines()
                has_real_prose = any(len(ln.split()) > 5 for ln in lines)

                if has_real_prose:
                    cleaned = "\n".join(
                        ln for ln in lines
                        if ln.strip().lower() not in (n_ch, n_sub)
                        and not re.match(r"^(chapter|day)\s+\d+:", ln.strip(), flags=re.I)
                    )

                # Hard guard: treat too-short/empty as failure so we retry
                if len(cleaned.split()) < MIN_HARD_WORDS:
                    cleaned = ""

                if cleaned and len(cleaned.split()) >= MIN_SUBSECTION_WORDS:
                    out_text = cleaned
                    break

                tries += 1

            # Final fallback: if still short, force a second-pass prompt that
            # explicitly expands with concrete examples/checklists/scripts (no headings)
            if not out_text:
                force_prompt = f"""{subheading_prompt(MASTER_PROMPT, persona, title, chapters_list, ch["title"], sub, mem)}
IMPORTANT:
- Your previous attempt was too short. Produce a FULL 700–800 words now.
- If content overlaps previous sections, add NEW case examples, therapist language, checklists or scripts.
- Absolutely NO headings; continuous prose only."""
                raw = call_openai(force_prompt, SUBSECTION_TOKENS)
                cleaned = clean_text(raw)
                if len(cleaned.split()) < MIN_SUBSECTION_WORDS:
                    # As a last resort, accept what we have but we still write it so nothing is "skipped"
                    out_text = cleaned
                else:
                    out_text = cleaned

            fixed = quick_validate_and_fix(out_text, mem.get("claims"), angle=mem.get("angle"))
            write_subsection(doc, fixed)

            mem = update_memory_from_text(fixed, mem)
            save_doc(doc, doc_path)

    save_doc(doc, doc_path)
    print(f"\n✅ Done. File generated:\n{doc_path.resolve()}\n")

if __name__ == "__main__":
    if not os.getenv("OPENAI_API_KEY"):
        print("Errore: manca OPENAI_API_KEY nell'ambiente.")
        sys.exit(1)
    main()
